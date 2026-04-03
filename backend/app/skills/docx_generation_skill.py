"""
DocxGenerationSkill - Word 文档生成

将错题列表导出为格式化的 Word 文档（.docx），
供学生打印或离线复习使用。
"""

import uuid
from pathlib import Path
from datetime import datetime

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from app.skills.base_skill import BaseSkill
from app.config import BACKEND_DIR
from app.utils.logging import get_logger

logger = get_logger("skill.docx_generation")

# 科目代码 → 中文名映射
SUBJECT_NAMES = {
    "ds": "数据结构",
    "os": "操作系统",
    "co": "计算机组成原理",
    "cn": "计算机网络",
}

# 导出文件存放目录
EXPORT_DIR = BACKEND_DIR / "data" / "exports"


class DocxGenerationSkill(BaseSkill):
    name = "docx_generation"
    description = "将错题导出为 Word 文档"

    def __init__(self):
        super().__init__()
        EXPORT_DIR.mkdir(parents=True, exist_ok=True)

    def _execute_impl(self, params: dict) -> dict:
        """
        生成错题 Word 文档。

        Args:
            params: {
                "mistakes": list[dict] - 错题列表，每项包含:
                    mistake_id, subject_code, page, chapter,
                    question_number, question_text, answer_text,
                    explanation, added_at
                "title": str - 文档标题 (可选，默认 "408考研错题集")
            }

        Returns:
            {"filename": str, "filepath": str, "download_url": str}
        """
        mistakes = params.get("mistakes", [])
        title = params.get("title", "408考研错题集")

        if not mistakes:
            self.logger.warning("无错题数据，跳过生成")
            return {
                "filename": "",
                "filepath": "",
                "download_url": "",
                "count": 0,
            }

        self.logger.info("开始生成 Word 文档 title='%s' count=%d", title, len(mistakes))

        doc = Document()
        self._setup_styles(doc)
        self._add_title(doc, title)
        self._add_summary(doc, mistakes)

        # 按科目分组
        by_subject = self._group_by_subject(mistakes)

        for subject_code, subject_mistakes in by_subject.items():
            subject_name = SUBJECT_NAMES.get(subject_code, subject_code)
            self._add_subject_section(doc, subject_name, subject_mistakes)

        # 保存文件
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        short_id = uuid.uuid4().hex[:6]
        filename = f"错题集_{timestamp}_{short_id}.docx"
        filepath = EXPORT_DIR / filename

        doc.save(str(filepath))

        self.logger.info("Word 文档生成完成 filename=%s size=%d bytes", filename, filepath.stat().st_size)

        return {
            "filename": filename,
            "filepath": str(filepath),
            "download_url": f"/api/mistakes/download/{filename}",
            "count": len(mistakes),
        }

    def _setup_styles(self, doc: Document) -> None:
        """配置文档默认样式"""
        style = doc.styles["Normal"]
        font = style.font
        font.name = "微软雅黑"
        font.size = Pt(10.5)

        # 设置页边距
        for section in doc.sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)

    def _add_title(self, doc: Document, title: str) -> None:
        """添加文档标题"""
        heading = doc.add_heading(title, level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 日期信息
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = date_para.add_run(f"生成日期: {datetime.now().strftime('%Y年%m月%d日')}")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(128, 128, 128)

    def _add_summary(self, doc: Document, mistakes: list[dict]) -> None:
        """添加统计摘要"""
        total = len(mistakes)
        by_subject = self._group_by_subject(mistakes)

        summary_text = f"共计 {total} 道错题"
        parts = []
        for code, items in by_subject.items():
            name = SUBJECT_NAMES.get(code, code)
            parts.append(f"{name} {len(items)} 题")
        if parts:
            summary_text += "（" + "、".join(parts) + "）"

        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(summary_text)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(80, 80, 80)

        doc.add_paragraph("")  # 空行

    def _add_subject_section(
        self, doc: Document, subject_name: str, mistakes: list[dict]
    ) -> None:
        """添加科目分节"""
        doc.add_heading(f"{subject_name}（{len(mistakes)} 题）", level=1)

        # 按章节排序
        sorted_mistakes = sorted(
            mistakes,
            key=lambda m: (m.get("chapter", ""), m.get("question_number", 0)),
        )

        for idx, mistake in enumerate(sorted_mistakes, 1):
            self._add_mistake_entry(doc, idx, mistake)

    def _add_mistake_entry(self, doc: Document, idx: int, mistake: dict) -> None:
        """添加单条错题"""
        chapter = mistake.get("chapter", "")
        page = mistake.get("page", 0)
        q_num = mistake.get("question_number", 0)
        question_text = mistake.get("question_text", "")
        answer_text = mistake.get("answer_text", "")
        explanation = mistake.get("explanation", "")
        added_at = mistake.get("added_at", "")

        # 题目标题
        heading_text = f"第{idx}题 — 第{chapter}章 P{page} 第{q_num}题"
        heading = doc.add_heading(heading_text, level=2)

        # 添加日期标签
        if added_at:
            para = doc.add_paragraph()
            run = para.add_run(f"添加时间: {added_at}")
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(160, 160, 160)

        # 题目内容
        if question_text:
            doc.add_heading("题目", level=3)
            para = doc.add_paragraph(question_text)
            para.paragraph_format.left_indent = Cm(0.5)

        # 答案
        if answer_text:
            doc.add_heading("答案", level=3)
            para = doc.add_paragraph(answer_text)
            para.paragraph_format.left_indent = Cm(0.5)

        # 解析
        if explanation:
            doc.add_heading("解析", level=3)
            para = doc.add_paragraph(explanation)
            para.paragraph_format.left_indent = Cm(0.5)

        # 分隔线
        doc.add_paragraph("─" * 50)

    @staticmethod
    def _group_by_subject(mistakes: list[dict]) -> dict[str, list[dict]]:
        """按科目分组"""
        grouped: dict[str, list[dict]] = {}
        for m in mistakes:
            code = m.get("subject_code", "unknown")
            grouped.setdefault(code, []).append(m)
        return grouped
